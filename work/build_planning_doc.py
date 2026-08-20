from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path

OUT = Path('outputs/Fitness_Coaching_Platform_Planning_Document.docx')
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = '16324F'; TEAL = '187A82'; PALE = 'EAF4F4'; LIGHT = 'F3F6F8'; MID = '5E6D78'; WHITE='FFFFFF'; RED='9B2C2C'; GOLD='9A6B16'

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(0.78)
sec.left_margin = sec.right_margin = Inches(0.86)
sec.header_distance = Inches(.35); sec.footer_distance = Inches(.35)

styles = doc.styles
normal = styles['Normal']; normal.font.name='Aptos'; normal.font.size=Pt(10.2); normal.font.color.rgb=RGBColor.from_string(NAVY)
normal.paragraph_format.space_after=Pt(5); normal.paragraph_format.line_spacing=1.08
for name,size,color,before,after in [('Title',30,NAVY,0,8),('Subtitle',13,MID,0,12),('Heading 1',17,NAVY,15,7),('Heading 2',12.5,TEAL,10,4),('Heading 3',10.5,NAVY,7,3)]:
    s=styles[name]; s.font.name='Aptos Display' if name!='Normal' else 'Aptos'; s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=name!='Subtitle'; s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)
def margins(cell, top=90, start=120, bottom=90, end=120):
    tc=cell._tc.get_or_add_tcPr(); m=tc.first_child_found_in('w:tcMar')
    if m is None: m=OxmlElement('w:tcMar'); tc.append(m)
    for tag,val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        x=m.find(qn('w:'+tag))
        if x is None: x=OxmlElement('w:'+tag); m.append(x)
        x.set(qn('w:w'),str(val)); x.set(qn('w:type'),'dxa')
def set_cell_text(cell, text, bold=False, color=NAVY, size=9):
    cell.text=''; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.02
    r=p.add_run(str(text)); r.bold=bold; r.font.name='Aptos'; r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(cell)
def table(headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,h in enumerate(headers):
        set_cell_text(t.rows[0].cells[i],h,True,WHITE,8.7); shade(t.rows[0].cells[i],NAVY)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,v in enumerate(row): set_cell_text(cells[i],v,False,NAVY,8.6); shade(cells[i], WHITE if ri%2==0 else LIGHT)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t
def bullet(text, level=0):
    p=doc.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2'); p.paragraph_format.space_after=Pt(3); p.add_run(text); return p
def new_numbering():
    numbering = doc.part.numbering_part.element
    nums = numbering.findall(qn('w:num'))
    num_id = max([int(n.get(qn('w:numId'))) for n in nums] + [0]) + 1
    # Reuse Word's built-in List Number abstract definition, but create a fresh
    # numbering instance so each logical sequence restarts at 1.
    base = None
    style_numpr = doc.styles['List Number']._element.find('.//' + qn('w:numPr'))
    style_numid = style_numpr.find(qn('w:numId')).get(qn('w:val')) if style_numpr is not None else None
    for n in nums:
        if n.get(qn('w:numId')) == style_numid and n.find(qn('w:abstractNumId')) is not None:
            base = n.find(qn('w:abstractNumId')).get(qn('w:val')); break
    if base is None: base = '0'
    num_el=OxmlElement('w:num'); num_el.set(qn('w:numId'),str(num_id))
    aid=OxmlElement('w:abstractNumId'); aid.set(qn('w:val'),base); num_el.append(aid)
    ov=OxmlElement('w:lvlOverride'); ov.set(qn('w:ilvl'),'0')
    st=OxmlElement('w:startOverride'); st.set(qn('w:val'),'1'); ov.append(st); num_el.append(ov)
    numbering.append(num_el)
    return num_id
def num(text, num_id):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3)
    pPr=p._p.get_or_add_pPr(); numPr=OxmlElement('w:numPr')
    ilvl=OxmlElement('w:ilvl'); ilvl.set(qn('w:val'),'0'); numPr.append(ilvl)
    nid=OxmlElement('w:numId'); nid.set(qn('w:val'),str(num_id)); numPr.append(nid); pPr.append(numPr)
    p.add_run(text); return p
def callout(label,text,fill=PALE):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; t.columns[0].width=Inches(6.65)
    c=t.cell(0,0); shade(c,fill); margins(c,150,170,150,170); c.text=''; p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0)
    r=p.add_run(label.upper()+'  '); r.bold=True; r.font.color.rgb=RGBColor.from_string(TEAL); r.font.size=Pt(9)
    r=p.add_run(text); r.font.color.rgb=RGBColor.from_string(NAVY); r.font.size=Pt(9.4)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
def section(title, intro=None):
    doc.add_heading(title, level=1)
    if intro: doc.add_paragraph(intro)
def h2(title): doc.add_heading(title,level=2)

# Header/footer
hp=sec.header.paragraphs[0]; hp.text='FITNESS COACHING PLATFORM  /  PRODUCT PLANNING'; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
for r in hp.runs: r.font.name='Aptos'; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(MID); r.bold=True
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=fp.add_run('CONFIDENTIAL WORKING DRAFT   •   '); r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(MID)
fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); fp._p.append(fld)

# Cover
doc.add_paragraph().paragraph_format.space_after=Pt(68)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(10)
r=p.add_run('PRODUCT PLANNING DOCUMENT'); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string(TEAL)
p=doc.add_paragraph(style='Title'); p.add_run('Fitness Coaching\nPlatform')
p=doc.add_paragraph(style='Subtitle'); p.add_run('Secure trainer–trainee collaboration, workout delivery, and progress tracking')
callout('Product thesis','Replace fragmented spreadsheets, messaging apps, and paper records with one shared workspace that is fast during workouts, clear about data ownership, and protective of sensitive health-related information.')
doc.add_paragraph().paragraph_format.space_after=Pt(70)
table(['DOCUMENT STATUS','RELEASE FOCUS','PRIMARY MARKET'],[['Working concept','Responsive web MVP','Independent and online trainers']], [2.15,2.15,2.15])
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Prepared for discovery, validation, architecture, and pilot planning'); r.italic=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string(MID)
doc.add_page_break()

section('Executive Summary')
doc.add_paragraph('The proposed platform connects fitness trainers and trainees in a shared, permission-aware workspace. Trainers manage clients, create reusable workout templates, assign programs, monitor completion, and provide nutrition guidance. Trainees view assigned plans, record results, log optional progress metrics, and control permitted personal information.')
table(['MVP OUTCOME','CORE USERS','DELIVERY SHAPE'],[
    ['Reliable coaching workflow from invitation through workout review','Trainers and trainees; organization administrators later','Responsive web application backed by a modular monolith'],
    ['Traceable progress history with stable workout records','Independent coaches, online coaches, and small teams','Spring Boot REST API, SvelteKit or React, PostgreSQL'],
    ['Sensitive-data safeguards and explicit access rules','Clients who need plans, metrics, guidance, and visibility','Containerized deployment, HTTPS, object storage for media']],[2.2,2.05,2.2])
callout('Recommended MVP boundary','Authentication, roles, profiles, trainer–trainee relationships, workouts, logging, basic progress and nutrition, dashboards, server-side permissions, and audit-friendly timestamps. Defer messaging, billing, wearables, native apps, food databases, and public discovery.')

section('1. Product Overview')
h2('Working concept'); doc.add_paragraph('A secure, mobile-friendly web platform where trainers manage client profiles, assign workouts, monitor progress, and provide nutrition guidance, while trainees view plans, record outcomes, and update permitted information.')
h2('Primary goal'); doc.add_paragraph('Consolidate coaching and progress tracking into one shared workspace, replacing fragmented spreadsheets, messaging applications, and paper records.')
h2('Target users')
for x in ['Trainers — personal trainers, independent coaches, online fitness coaches, and small training teams.','Trainees — clients accessing assigned workouts, nutrition guidance, progress metrics, and coaching communication.','Administrators — a future gym or organization role for managing trainers and clients.']: bullet(x)
h2('Product principles')
for x in ['Shared data has clear ownership and edit permissions.','The interface is fast enough for use during a workout.','History, summaries, and charts make progress understandable.','Health-related data is treated as sensitive.','The first release favors reliable core workflows over feature breadth.']: bullet(x)

section('2. Scope and Assumptions')
h2('Initial assumptions')
for x in ['A trainer manages multiple trainees.','A trainee works with one trainer initially; multi-trainer support may follow.','Trainers assign reusable workout templates.','Trainees record sets, repetitions, load, duration, and notes.','Weight, measurements, nutrition entries, and progress photos are optional and permission-controlled.','The product supports coaching and tracking; it is not a medical diagnostic system.']: bullet(x)
h2('Recommended MVP')
for x in ['Account registration, login, logout, email verification, and password reset.','Trainer and trainee roles with role-specific profiles.','Invitations and trainer–trainee relationship management.','Workout creation, assignment, scheduling, partial completion, and history.','Trainer-created exercise library.','Weight and generic body-metric logging.','Daily nutrition notes with optional calories and macronutrients.','Trainer and trainee dashboards with tables and basic charts.','Role-based permissions, audit-friendly timestamps, and responsive desktop/mobile web.']: bullet(x)
h2('Later releases')
doc.add_paragraph('Native apps; in-app messaging and notifications; calendar and appointments; structured meal planning and food databases; wearables; subscriptions and payments; organization accounts; exercise videos; automated insights; public trainer discovery and reviews.')

section('3. User Roles and Permissions')
table(['CAPABILITY','TRAINER','TRAINEE'],[
['Create/edit own profile','Yes','Yes'],['Invite/remove trainees','Yes','No'],['View connected profile','Yes','Yes'],['Create workout templates','Yes','No'],['Assign workouts','Yes','No'],['Edit assigned structure','Yes','Restricted / request-based'],['Log workout completion','Yes, if permitted','Yes'],['View workout history','Yes','Yes'],['Log weight/measurements','On behalf, if permitted','Yes'],['View nutrition records','Connected trainees','Own records'],['Edit trainer nutrition guidance','Yes','No'],['Add personal nutrition entries','Optional','Yes'],['Export personal data','Own data','Own data'],['Delete account','Yes','Yes']],[3.1,1.75,1.75])
callout('Enforcement rule','Permissions must be enforced at the API and database-access layers—not merely by hiding controls in the interface.')

section('4. Core User Stories')
h2('Trainer')
for x in ['Create a professional profile.','Invite a trainee by email or code.','Create reusable workout templates.','Assign a workout plan with start date and schedule.','See completion status and review results.','Record or review weight and nutrition progress.','Add coaching notes and guidance.','Archive a trainee without losing history.']: bullet('As a trainer, I want to '+x[0].lower()+x[1:])
h2('Trainee')
for x in ['View today’s assigned workout.','Record sets, reps, load, duration, and perceived difficulty.','Save partial progress, complete a workout, and add notes.','Update weight and nutrition entries.','View progress charts over time.','Understand trainer visibility and edit rights.','Remove trainer access when coaching ends.']: bullet('As a trainee, I want to '+x[0].lower()+x[1:])

section('5. Main Workflows')
for title, steps in [
('Trainer onboarding',['Register and verify email.','Select trainer role.','Complete profile: name, bio, specialties, optional credentials.','Create or adopt exercises.','Invite a trainee.']),
('Trainee onboarding',['Register from an invitation or independently.','Select trainee role.','Set optional profile, goals, units, and privacy preferences.','Accept trainer connection.','View dashboard and first workout.']),
('Workout assignment',['Trainer creates a versioned template.','Add ordered exercises and prescribed targets.','Assign to one or more trainees.','Set start date, frequency, and optional end date.','Trainee logs actual performance separately.','Trainer reviews completion, results, and notes.']),
('Progress tracking',['Select metric.','Enter value, date, unit, and note.','Validate and store author/timestamp.','Display table and chart history.','Permit authorized contextual or coaching notes.'])]:
    h2(title)
    sequence_id = new_numbering()
    for s in steps: num(s, sequence_id)

section('6. Functional Requirements')
reqs={
'Authentication and accounts':['Email/password initially; trusted strong one-way password hashing.','Email verification, password reset, session expiration, logout from other sessions.','Explicit onboarding role selection; role changes require a controlled process.'],
'Profiles':['Trainer fields: name, photo, bio, specialties, certifications, experience, location, contact preferences.','Trainee fields: name, photo, goals, units, experience level, and voluntarily supplied limitations. Collect date of birth or emergency information only with a defined need and safeguards.'],
'Relationships':['Invitations expire and require explicit acceptance.','Either party can pause or end a relationship.','Statuses: pending, active, paused, archived, revoked.','Historical access follows a documented retention policy.'],
'Exercise library':['Exercise definition may include muscle groups, equipment, instructions, difficulty, media reference, creator, visibility, and version.','Distinguish platform exercises from trainer-created exercises; protect exercises referenced by historical workouts.'],
'Workouts and programs':['Support names, descriptions, goals, duration, ordered exercises, set/rep/load/time/distance/rest/tempo targets, notes, and substitutions.','Lifecycle states: draft, published, assigned, completed, skipped, archived.','Use immutable versions or assignment snapshots.'],
'Workout logging':['Store actual performance separately from prescribed targets.','Support set-level reps, load, time, distance, rest, exertion, pain flag, and notes.','Allow partial completion and resumable saves.'],
'Progress metrics':['Use a generic metric model: type, value, unit, date/time, source, author, visibility, note.','Convert kilograms and pounds consistently while preserving original input where useful.'],
'Nutrition':['Daily/meal-level date, type, description, optional calories/macros/water.','Separate trainer guidance from trainee-reported intake.','Keep fields optional and avoid presenting guidance as medical advice.'],
'Dashboards':['Trainer: active trainees, invitations, upcoming assignments, completion, recent progress, attention items.','Trainee: today’s workout, current program, history, latest progress/nutrition, charts, trainer notes.']}
for k,items in reqs.items():
    h2(k)
    for x in items: bullet(x)

section('7. Suggested Technical Architecture')
table(['LAYER','RECOMMENDATION','RATIONALE'],[
['Frontend','SvelteKit or React','Responsive, mobile-first interaction'],['Backend','Spring Boot REST API','Strong Java structure, validation, and service boundaries'],['Database','PostgreSQL; SQLite for local prototype','Relational integrity and mature operations'],['Authentication','Secure sessions or short-lived access tokens with refresh rotation','Predictable session security'],['Files','Object storage','Keep large images outside relational rows'],['Charts','Frontend chart library','Weight, completion, nutrition trends'],['Deployment','Containers, managed PostgreSQL, HTTPS, environment configuration','Repeatable and secure operations']],[1.3,2.05,3.25])
h2('Logical components')
doc.add_paragraph('Identity and access; user/profile; trainer–trainee relationships; exercise/workout; workout logging; progress/nutrition; optional notifications; audit/privacy.')
callout('Architecture direction','Build an MVP modular monolith rather than microservices. Preserve clear module boundaries while minimizing operational complexity.')

section('8. Initial Data Model')
table(['ENTITY','PURPOSE'],[
['users','Identity, email, credential/provider, role, status, timestamps'],['trainer_profiles / trainee_profiles','Role-specific profile data linked to users'],['trainer_trainee_relationships','Parties, status, invitation, permissions, timestamps'],['exercises','Definition, creator, visibility, equipment, version'],['workout_templates / template_exercises','Reusable definitions and ordered targets'],['assigned_workouts','Template snapshot/version, trainee, schedule, status'],['workout_logs / set_logs','Completion metadata and actual set-level performance'],['progress_metrics / progress_entries','Metric definitions and values with author, unit, visibility'],['nutrition_entries','Daily or meal-level reported nutrition'],['trainer_notes','Private or shared coaching notes'],['notifications','Recipient, event, read state, timestamps'],['audit_events','Actor, action, entity, timestamp, limited request metadata']],[2.35,4.25])
h2('Modeling rules')
for x in ['Use UUIDs or comparable non-sequential public identifiers.','Include created_at, updated_at, and relevant deleted_at timestamps.','Use database constraints for relationship validity, nonnegative values, and unique active connections.','Store units explicitly.','Use soft deletion where auditability or historical accuracy requires it.']: bullet(x)

section('9. API Outline')
table(['DOMAIN','ENDPOINTS'],[
['Authentication','POST /api/auth/register • login • logout • forgot-password • reset-password'],['Profiles & relationships','GET /api/me • PATCH /api/me/profile • POST /api/trainers/invitations • GET/PATCH relationships'],['Exercises & workouts','GET/POST /api/exercises • POST/GET workout-templates • POST assigned-workouts • GET trainee workouts • POST workout logs'],['Progress & nutrition','GET/POST/PATCH progress-entries • GET/POST/PATCH nutrition-entries']],[1.7,4.9])
doc.add_paragraph('Every endpoint enforces ownership or active-relationship authorization, validates server-side payloads, and returns a consistent error envelope.')

section('10. Privacy and Security')
for x in ['Collect only information required by a defined feature.','Use HTTPS outside local development.','Use modern password hashing; never log credentials or tokens.','Apply role- and relationship-based authorization to every protected resource.','Test broken object-level authorization, especially URL identifiers.','Use secure cookies or carefully designed token storage, CSRF controls where applicable, and authentication rate limits.','Encrypt backups and tightly restrict production database access.','Keep private notes, nutrition, and photos private by default.','Provide deletion, export, consent controls, and a clear privacy policy.','Audit sensitive actions without unnecessary personal content.','Define retention for inactive accounts, invitations, logs, and media.','Review Canadian privacy obligations with qualified legal counsel before launch.']: bullet(x)
callout('Security priority','Authorization failures are the highest-impact product risk. Design object-level access checks as domain rules and cover them with integration and API tests.', 'FCECEC')

section('11. Non-Functional Requirements')
table(['QUALITY','REQUIREMENT'],[
['Performance','Common dashboard and workout screens load quickly on mobile connections.'],['Availability','Charts, notifications, and uploads fail gracefully without blocking core flows.'],['Accessibility','Keyboard support, readable contrast, form labels, screen-reader status messages.'],['Responsiveness','Core logging works on small screens with large touch targets.'],['Reliability','Workout save is idempotent or guarded against duplicate submission.'],['Observability','Structured logs, error tracking, health checks, privacy-safe usage metrics.'],['Maintainability','Domain logic outside controllers; migrations, API docs, consistent errors and validation.']],[1.45,5.15])

section('12. MVP Delivery Plan')
table(['PHASE','FOCUS','EXIT SIGNAL'],[
['1. Discovery & design','Customer segment, scope, flows, wireframes, nutrition model, permissions, retention','Validated MVP boundary and testable flows'],['2. Foundation','Repositories, CI, containers, config, migrations, auth, roles, profiles, tests','Secure account and role workflows'],['3. Relationships','Invitations, acceptance, status changes, revocation, dashboards','Connected trainer–trainee workflow'],['4. Workouts','Exercises, templates, assignment, scheduling, logging, snapshots','Mobile workout flow usable end to end'],['5. Progress & nutrition','Generic metrics, journal, macros, charts, history export','Trends visible to authorized parties'],['6. Hardening & release','Abuse cases, rate limits, backups, monitoring, production, pilot','Controlled pilot ready']],[1.15,3.35,2.1])

section('13. Testing Strategy')
table(['TEST LAYER','FOCUS'],[
['Unit','Workout calculations, validation, conversion, permissions, domain rules'],['Integration','Database relationships, invitations, assignments, entry access'],['API','Authentication, authorization, malformed input, duplicates, forbidden access'],['End-to-end','Trainer onboarding, acceptance, completion, progress review'],['Security','ID tampering, escalation, sessions, rate limits, injection, uploads'],['Usability','One trainer builds and assigns; one trainee logs on a phone unaided']],[1.45,5.15])

section('14. Success Metrics')
for x in ['Invite-to-activation conversion.','Time to first created and assigned workout.','Assigned workouts opened and completed.','Active trainees logging progress weekly.','Trainer and trainee weekly active users.','Workout-log save failure rate.','Permission or missing-data support issues.','30-day and 90-day retention.']: bullet(x)

section('15. Risks and Mitigations')
table(['RISK','MITIGATION'],[
['Trainer setup takes too long','Defaults, reusable templates, duplication, bulk assignment'],['Users misunderstand edit rights','Field-level ownership and permission indicators'],['History changes unexpectedly','Snapshot or version assigned workout data'],['Sensitive data exposure','Server authorization, data minimization, object-access testing'],['Nutrition scope expands','Journal plus optional macros; defer food databases'],['Mobile logging is awkward','Mobile-first logger with minimal typing'],['Scope expands rapidly','Written MVP boundary and validated-problem prioritization'],['Coaching styles differ','Optional metrics, notes, and configurable fields']],[2.3,4.3])

section('16. Recommended First Backlog')
backlog_id = new_numbering()
for x in ['Define personas, MVP boundaries, and permission matrix.','Wireframe trainer dashboard, trainee dashboard, workout builder, workout logger, and progress screen.','Set up Spring Boot, frontend, PostgreSQL, migrations, containers, and CI.','Implement registration, login, roles, and profiles.','Implement invitations and relationship state changes.','Implement exercise library and versioned workout templates.','Implement assignment and trainee logging.','Implement progress and nutrition entries.','Add dashboards, charts, filters, and exports.','Add authorization tests, audit events, backups, and deployment controls.','Pilot with a small trainer/trainee cohort.','Use pilot evidence to select the next release.']: num(x, backlog_id)

section('17. Definition of MVP Done')
for x in ['A trainer can register, create a profile, invite a trainee, create a workout, and assign it.','A trainee can accept, view the assignment, save partial or complete results, and add notes.','Both parties can view authorized weight and nutrition history.','Changing a URL identifier never exposes another user’s records.','Historical assignments remain stable after template edits.','Core flows work in a mobile browser and critical authorization paths have automated coverage.','A privacy-policy draft, backup process, error monitoring, and documented deletion process exist.']: bullet(x)

section('18. Product Decisions to Make Next')
decisions=['Responsive web only for release one, or a native app?','Who pays: trainers, trainees, or both?','Can a trainee work with multiple trainers?','Which fields can trainees edit, and can trainers override them?','Free-form nutrition journal or structured macro tracking?','Are progress photos in the MVP?','Is in-platform communication required?','Which data must be exportable or deletable?','Which Canadian privacy and business requirements apply?','What is the smallest credible pilot group?']
table(['#','DECISION','WORKING DIRECTION'],[(i+1,d,'Confirm in discovery') for i,d in enumerate(decisions)],[.45,4.55,1.6])
callout('Immediate next step','Run a short discovery cycle with 5–8 independent trainers and representative trainees. Validate the mobile workout logger, invitation flow, permission expectations, nutrition depth, and willingness to pay before committing to implementation scope.')

# Keep tables together where practical and repeat header rows.
for t in doc.tables:
    t.rows[0]._tr.get_or_add_trPr().append(OxmlElement('w:tblHeader'))
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs: p.paragraph_format.keep_together=True

doc.core_properties.title='Fitness Coaching Platform Planning Document'
doc.core_properties.subject='MVP product, architecture, privacy, delivery, and pilot plan'
doc.core_properties.author='Product Planning'
doc.save(OUT)
print(OUT.resolve())
